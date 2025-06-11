from xml.etree import ElementTree as ET
from docx import Document
from docx.shared import Pt


nessus_file_path = r"/PATH/TO/.nessus_file"


tree = ET.parse(nessus_file_path)
root = tree.getroot()


doc = Document()


severity_levels = {
    "4": "Critical",
    "3": "High",
    "2": "Medium",
    "1": "Low"
}


vulnerabilities = {level: [] for level in severity_levels.values()}

#Vul. extract --> nessus.db 

for report_item in root.findall(".//ReportItem"):
    severity = report_item.get("severity")
    if severity in severity_levels:
        vuln_data = {
            "Name": report_item.get("pluginName", "N/A"),
            "Criticality": severity_levels[severity],
            "Synopsis": report_item.findtext("synopsis", "N/A"),
            "Description": report_item.findtext("description", "N/A"),
            "Solution": report_item.findtext("solution", "N/A"),
            "Plugin Output": report_item.findtext("plugin_output", "N/A")
        }
        vulnerabilities[severity_levels[severity]].append(vuln_data)


for severity, vulns in vulnerabilities.items():
    if vulns:
        doc.add_heading(f"{severity} Vulnerabilities", level=1)
        for vuln in vulns:
           
            doc.add_heading(vuln["Name"], level=2)
            
           
            table = doc.add_table(rows=5, cols=2)
            table.style = 'Table Grid'

            
            headings = ["Criticality", "Synopsis", "Description", "Solution", "Plugin Output"]

          
            for idx, heading in enumerate(headings):
                row = table.rows[idx]
                row.cells[0].text = heading
                row.cells[1].text = vuln.get(heading, "N/A")

               
                cell = row.cells[0]
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.font.size = Pt(10)
                run.font.name = 'Arial'

            doc.add_paragraph()  


output_file_path = r"/PATH/TO/Save_File"
doc.save(output_file_path)
